from fastapi import UploadFile
from typing import List, Optional, Union
import io
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom

async def convert_uploadfile_to_pil_image(upload_file: UploadFile) -> Image.Image:
    # Чтение байтов из файла
    contents = await upload_file.read()
    # Создание объекта BytesIO из байтов
    image_stream = io.BytesIO(contents)

    pil_image = Image.open(image_stream)
    pil_image.load()
    
    return pil_image


def convert_pdf_to_images(
    pdf_bytes: bytes,
    dpi: int = 200,
    first_page: Optional[int] = None,
    last_page: Optional[int] = None,
    fmt: str = 'jpeg',
    thread_count: int = 4
) -> List[Image.Image]:
    try:
        images = convert_from_bytes(
            pdf_bytes,
            dpi=dpi,
            first_page=first_page,
            last_page=last_page,
            fmt=fmt,
            thread_count=thread_count,
            use_pdftocairo=True,
            transparent=False
        )
        
        # Конвертируем в RGB, если нужно (для JPEG)
        if fmt.lower() == 'jpeg':
            images = [image.convert('RGB') for image in images]
            
        return images
        
    except Exception as e:
        raise ValueError(f"Ошибка конвертации PDF: {str(e)}")

def get_text_from_docx(doc_bytes: bytes) -> str:
    root = ET.Element("extracted_text")
    
    # Открываем документ из байтов
    doc_stream = io.BytesIO(doc_bytes)
    try:
        doc = Document(doc_stream)
        
        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Проверяем, что параграф не пустой
                text_elem = ET.SubElement(root, "text")
                text_elem.text = paragraph.text.strip()
                # Fixed: Changed print.paragraph to print(paragraph.text)
        
        # Обрабатываем таблицы
        for table in doc.tables:
            table_elem = ET.SubElement(root, "table")
            
            for row in table.rows:
                row_elem = ET.SubElement(table_elem, "row")
                
                for cell in row.cells:
                    # Используем "coloumn" как в примере (хотя правильно было бы "column")
                    col_elem = ET.SubElement(row_elem, "coloumn")
                    # Извлекаем текст из ячейки
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text += paragraph.text.strip() + " "
                    col_elem.text = cell_text.strip()
        
        # Преобразуем XML в строку с красивым форматированием
        rough_string = ET.tostring(root, 'utf-8')
        reparsed = minidom.parseString(rough_string)
        return reparsed.toprettyxml(indent="  ").replace('<?xml version="1.0" ?>\n', '')
        
    except Exception as e:
        # Если произошла ошибка, возвращаем XML с сообщением об ошибке
        error_elem = ET.SubElement(root, "text")
        error_elem.text = f"Ошибка при извлечении текста: {str(e)}"
        rough_string = ET.tostring(root, 'utf-8')
        reparsed = minidom.parseString(rough_string)
        return reparsed.toprettyxml(indent="  ").replace('<?xml version="1.0" ?>\n', '')


def get_png_payload(prompt: str, images: List[Image.Image]) -> List[Union[str, Image.Image]]:
    contents = [prompt]
    contents.extend(images)
    return contents

def get_pdf_payload(prompt: str, pdf_bytes: bytes) -> list:
    contents = [prompt, pdf_bytes]
    
    return contents
